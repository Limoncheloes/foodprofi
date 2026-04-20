"""Document generation service for procurement orders.

Generates:
- DOCX (python-docx): printable order sheet grouped by buyer
- XLSX (openpyxl): 1C import format, two sheets

Both functions accept plain data objects (dicts or mock-friendly objects)
so they can be unit-tested without a database.
"""
import io
from typing import Any

from docx import Document
from docx.shared import Cm, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import openpyxl
from openpyxl.styles import Font, PatternFill, Alignment


# Excel column headers — update to match client's 1C import template
XLSX_COLUMNS = [
    "Дата",
    "Номер заявки",
    "Ресторан",
    "Наименование",
    "Артикул",
    "Кол-во заказано",
    "Кол-во получено",
    "Ед. изм.",
    "Закупщик",
    "Категория",
    "Примечание",
]


def generate_docx(order: Any, items: list[Any]) -> bytes:
    """Generate a printable Word document grouped by buyer.

    order attributes used: id, restaurant_name, user_name, created_at
    item attributes used: display_name, quantity_ordered, quantity_received,
                          unit, buyer_name, category_name
    """
    doc = Document()

    # Page margins: 2cm all sides
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2)
        section.right_margin = Cm(2)

    # Title
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run(f"ЗАЯВКА НА ЗАКУПКУ №{str(order.id)[:8].upper()}")
    run.bold = True
    run.font.size = Pt(14)

    # Header info
    date_str = order.created_at.strftime("%d.%m.%Y")
    info = doc.add_paragraph()
    info.add_run(
        f"Дата: {date_str}  |  Ресторан: {order.restaurant_name}  |  Создал: {order.user_name}"
    ).font.size = Pt(11)

    doc.add_paragraph()  # spacer

    # Group items by buyer
    buyers: dict[str, list[Any]] = {}
    for item in items:
        buyer_name = getattr(item, "buyer_name", None) or "Не назначен"
        buyers.setdefault(buyer_name, []).append(item)

    for buyer_name, buyer_items in buyers.items():
        # Buyer separator
        sep = doc.add_paragraph()
        sep.add_run("─" * 60)

        buyer_header = doc.add_paragraph()
        run = buyer_header.add_run(f"ЗАКУПЩИК: {buyer_name}")
        run.bold = True
        run.font.size = Pt(12)

        sep2 = doc.add_paragraph()
        sep2.add_run("─" * 60)

        # Table
        table = doc.add_table(rows=1, cols=5)
        table.style = "Table Grid"
        hdr = table.rows[0].cells
        for cell, text in zip(hdr, ["№", "Наименование", "Заказано", "Получено", "Ед."]):
            hdr_run = cell.paragraphs[0].add_run(text)
            hdr_run.bold = True
            hdr_run.font.size = Pt(11)

        for idx, item in enumerate(buyer_items, 1):
            row = table.add_row().cells
            qty_ordered = f"{float(item.quantity_ordered):.3f}"
            qty_received = "_______" if item.quantity_received is None else f"{float(item.quantity_received):.3f}"
            for cell, text in zip(row, [str(idx), item.display_name, qty_ordered, qty_received, item.unit]):
                cell.paragraphs[0].add_run(text).font.size = Pt(12)

        doc.add_paragraph()  # spacer between buyers

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def generate_xlsx(order: Any, items: list[Any]) -> bytes:
    """Generate an Excel file for 1C import.

    Sheet 1: all items (one row per item)
    Sheet 2: summary aggregated by category

    order attributes: id, restaurant_name, created_at
    item attributes: display_name, quantity_ordered, quantity_received,
                     unit, buyer_name, category_name, substitution_note
    """
    wb = openpyxl.Workbook()

    # ── Sheet 1: All items ──
    ws1 = wb.active
    ws1.title = "Все позиции"

    header_font = Font(bold=True, size=11)
    header_fill = PatternFill("solid", fgColor="D9D9D9")

    for col_idx, col_name in enumerate(XLSX_COLUMNS, 1):
        cell = ws1.cell(row=1, column=col_idx, value=col_name)
        cell.font = header_font
        cell.fill = header_fill
        cell.alignment = Alignment(horizontal="center")

    date_str = order.created_at.strftime("%d.%m.%Y")
    order_short = str(order.id)[:8].upper()

    for row_idx, item in enumerate(items, 2):
        note = getattr(item, "substitution_note", None) or ""
        ws1.cell(row=row_idx, column=1, value=date_str)
        ws1.cell(row=row_idx, column=2, value=order_short)
        ws1.cell(row=row_idx, column=3, value=order.restaurant_name)
        ws1.cell(row=row_idx, column=4, value=item.display_name)
        ws1.cell(row=row_idx, column=5, value="")  # артикул — пусто для некаталожных
        ws1.cell(row=row_idx, column=6, value=float(item.quantity_ordered))
        ws1.cell(row=row_idx, column=7, value=float(item.quantity_received) if item.quantity_received is not None else None)
        ws1.cell(row=row_idx, column=8, value=item.unit)
        ws1.cell(row=row_idx, column=9, value=getattr(item, "buyer_name", ""))
        ws1.cell(row=row_idx, column=10, value=getattr(item, "category_name", ""))
        ws1.cell(row=row_idx, column=11, value=note)

    # Auto-width columns
    for col in ws1.columns:
        max_len = max((len(str(c.value or "")) for c in col), default=0)
        ws1.column_dimensions[col[0].column_letter].width = min(max_len + 2, 40)

    # ── Sheet 2: Summary by category ──
    ws2 = wb.create_sheet("По категориям")
    summary_headers = ["Категория", "Наименование", "Итого заказано", "Итого получено", "Ед. изм."]
    for col_idx, col_name in enumerate(summary_headers, 1):
        cell = ws2.cell(row=1, column=col_idx, value=col_name)
        cell.font = header_font
        cell.fill = header_fill

    for row_idx, item in enumerate(items, 2):
        ws2.cell(row=row_idx, column=1, value=getattr(item, "category_name", "") or "Без категории")
        ws2.cell(row=row_idx, column=2, value=item.display_name)
        ws2.cell(row=row_idx, column=3, value=float(item.quantity_ordered))
        ws2.cell(row=row_idx, column=4, value=float(item.quantity_received) if item.quantity_received is not None else None)
        ws2.cell(row=row_idx, column=5, value=item.unit)

    buf = io.BytesIO()
    wb.save(buf)
    return buf.getvalue()
